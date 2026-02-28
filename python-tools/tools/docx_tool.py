"""DOCX read/write tools using python-docx."""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches


def docx_read(file_path: str) -> str:
    """Read a DOCX file and return its full text.

    Args:
        file_path: Path to the .docx file.

    Returns:
        The document text with paragraphs joined by newlines.
    """
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs)


def docx_write(
    file_path: str,
    content: str,
    template: str | None = None,
) -> str:
    """Create or overwrite a DOCX file.

    Args:
        file_path: Destination .docx path.
        content: Text content. Each line becomes a paragraph.
        template: Optional path to a .docx template to use as base.

    Returns:
        The path of the written file.
    """
    path = Path(file_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(template) if template else Document()

    for line in content.split("\n"):
        doc.add_paragraph(line)

    doc.save(path)
    return str(path)


def manuscript_generate(
    excel_path: str | None = None,
    figures: list[str] | None = None,
    sections: dict[str, str] | None = None,
    template: str | None = None,
    journal_style: str | None = None,
    output_path: str = "manuscript.docx",
) -> str:
    """Generate a manuscript draft DOCX.

    Args:
        excel_path: Path to data Excel file (for auto-generating tables).
        figures: List of figure image paths to embed.
        sections: Dict of section_name → content text.
        template: Optional .docx template path.
        journal_style: Optional journal style name (for heading convention).
        output_path: Destination .docx path.

    Returns:
        The path of the generated manuscript.
    """
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(template) if template else Document()

    # Title
    style = journal_style or "default"
    doc.add_heading(f"Manuscript Draft ({style})", level=0)

    # Sections
    default_sections = ["Abstract", "Introduction", "Methods", "Results", "Discussion", "Conclusion"]
    sec_content = sections or {}

    for sec_name in default_sections:
        doc.add_heading(sec_name, level=1)
        text = sec_content.get(sec_name, f"[{sec_name} content to be added]")
        doc.add_paragraph(text)

    # Data table from Excel
    if excel_path and Path(excel_path).exists():
        doc.add_heading("Data Summary", level=1)
        from tools.excel import excel_read

        rows = excel_read(excel_path)
        if rows:
            headers = list(rows[0].keys())
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = str(h)
            for r_idx, row in enumerate(rows, 1):
                for c_idx, h in enumerate(headers):
                    table.rows[r_idx].cells[c_idx].text = str(row.get(h, ""))

    # Figures
    if figures:
        doc.add_heading("Figures", level=1)
        for i, fig_path in enumerate(figures, 1):
            if Path(fig_path).exists():
                doc.add_paragraph(f"Figure {i}")
                doc.add_picture(fig_path, width=Inches(5))

    doc.save(path)
    return str(path)
